"""Tests for template_registry and template_service modules."""
import io
import pytest
from unittest.mock import MagicMock, patch

from app.template_registry import (
    TEMPLATE_REGISTRY,
    get_alternate_s3_keys,
    get_placeholder_map,
    get_template_mapping,
    get_template_s3_key,
    has_template,
    is_markdown_only,
    list_registered_doc_types,
)
from app.template_service import (
    DOCXPopulator,
    TemplateResult,
    TemplateService,
    XLSXPopulator,
    _cache_get,
    _cache_set,
    _MISS,
)


# ══════════════════════════════════════════════════════════════════════
#  Template Registry Tests
# ══════════════════════════════════════════════════════════════════════


class TestTemplateRegistry:
    """Tests for template_registry module."""

    def test_has_template_for_registered_types(self):
        """Registered doc_types should return True."""
        assert has_template("sow") is True
        assert has_template("igce") is True
        assert has_template("market_research") is True
        assert has_template("justification") is True
        assert has_template("acquisition_plan") is True
        assert has_template("cor_certification") is True

    def test_has_template_for_markdown_only(self):
        """Markdown-only types should not have templates."""
        assert has_template("eval_criteria") is False
        assert has_template("security_checklist") is False
        assert has_template("section_508") is False
        assert has_template("contract_type_justification") is False

    def test_is_markdown_only(self):
        """Markdown-only types should be identified correctly."""
        assert is_markdown_only("eval_criteria") is True
        assert is_markdown_only("security_checklist") is True
        assert is_markdown_only("sow") is False
        assert is_markdown_only("igce") is False

    def test_get_template_mapping_returns_mapping(self):
        """Should return TemplateMapping for registered types."""
        mapping = get_template_mapping("sow")
        assert mapping is not None
        assert mapping.doc_type == "sow"
        assert mapping.file_type == "docx"
        assert mapping.s3_filename.endswith(".docx")

    def test_get_template_mapping_returns_none_for_unknown(self):
        """Should return None for unregistered types."""
        assert get_template_mapping("unknown_type") is None
        assert get_template_mapping("eval_criteria") is None

    def test_get_template_s3_key(self):
        """Should return full S3 key for registered types."""
        key = get_template_s3_key("sow")
        assert key is not None
        assert "approved/supervisor-core/essential-templates" in key
        assert key.endswith(".docx")

    def test_get_template_s3_key_none_for_unknown(self):
        """Should return None for unregistered types."""
        assert get_template_s3_key("unknown_type") is None

    def test_get_alternate_s3_keys(self):
        """Should return alternate keys for types with alternates."""
        alts = get_alternate_s3_keys("igce")
        assert isinstance(alts, list)
        assert len(alts) >= 1  # IGCE has Educational and Nonprofit alternates

    def test_get_placeholder_map(self):
        """Should return placeholder mapping for registered types."""
        placeholders = get_placeholder_map("sow")
        assert "title" in placeholders
        assert placeholders["title"] == "{{PROJECT_TITLE}}"

    def test_list_registered_doc_types(self):
        """Should list all registered doc types."""
        types = list_registered_doc_types()
        assert "sow" in types
        assert "igce" in types
        assert len(types) == len(TEMPLATE_REGISTRY)


# ══════════════════════════════════════════════════════════════════════
#  Template Service Cache Tests
# ══════════════════════════════════════════════════════════════════════


class TestTemplateCache:
    """Tests for template caching."""

    def test_cache_miss_returns_sentinel(self):
        """Uncached items should return _MISS sentinel."""
        result = _cache_get("test-bucket", "nonexistent-key")
        assert result is _MISS

    def test_cache_set_and_get(self):
        """Cached items should be retrievable."""
        _cache_set("test-bucket", "test-key", b"test-data")
        result = _cache_get("test-bucket", "test-key")
        assert result == b"test-data"

    def test_cache_can_store_none(self):
        """Cache should distinguish between None (cached miss) and _MISS."""
        _cache_set("test-bucket", "none-key", None)
        result = _cache_get("test-bucket", "none-key")
        assert result is None  # Not _MISS


# ══════════════════════════════════════════════════════════════════════
#  DOCX Populator Tests
# ══════════════════════════════════════════════════════════════════════


class TestDOCXPopulator:
    """Tests for DOCXPopulator."""

    @pytest.fixture
    def sample_docx(self):
        """Create a minimal DOCX for testing."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Project: {{PROJECT_TITLE}}")
        doc.add_paragraph("Description: {{DESCRIPTION}}")
        doc.add_paragraph("Tasks: {{TASKS}}")

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    def test_populate_replaces_placeholders(self, sample_docx):
        """Placeholders should be replaced with data values."""
        data = {
            "title": "Cloud Migration",
            "description": "Migrate to AWS",
        }
        placeholder_map = {
            "title": "{{PROJECT_TITLE}}",
            "description": "{{DESCRIPTION}}",
        }

        result = DOCXPopulator.populate(sample_docx, data, placeholder_map)
        preview = DOCXPopulator.extract_text(result)

        assert "Cloud Migration" in preview
        assert "Migrate to AWS" in preview
        assert "{{PROJECT_TITLE}}" not in preview

    def test_populate_handles_lists(self, sample_docx):
        """List values should be converted to bullet points."""
        data = {
            "tasks": ["Task 1", "Task 2", "Task 3"],
        }
        placeholder_map = {
            "tasks": "{{TASKS}}",
        }

        result = DOCXPopulator.populate(sample_docx, data, placeholder_map)
        preview = DOCXPopulator.extract_text(result)

        assert "Task 1" in preview or "- Task 1" in preview

    def test_extract_text_returns_markdown(self, sample_docx):
        """Text extraction should return readable content."""
        preview = DOCXPopulator.extract_text(sample_docx)
        assert "Project:" in preview
        assert "Description:" in preview


# ══════════════════════════════════════════════════════════════════════
#  XLSX Populator Tests
# ══════════════════════════════════════════════════════════════════════


class TestXLSXPopulator:
    """Tests for XLSXPopulator."""

    @pytest.fixture
    def sample_xlsx(self):
        """Create a minimal XLSX for testing."""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        wb = Workbook()
        ws = wb.active
        ws.title = "IGCE"
        ws["A1"] = "Project: {{PROJECT_TITLE}}"
        ws["A2"] = "Total: {{TOTAL_ESTIMATE}}"
        ws["A3"] = "{{LINE_ITEMS}}"

        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()

    def test_populate_replaces_placeholders(self, sample_xlsx):
        """Placeholders should be replaced with data values."""
        data = {
            "title": "Server Procurement",
            "total_estimate": "$150,000",
        }
        placeholder_map = {
            "title": "{{PROJECT_TITLE}}",
            "total_estimate": "{{TOTAL_ESTIMATE}}",
        }

        result = XLSXPopulator.populate(sample_xlsx, data, placeholder_map)
        preview = XLSXPopulator.extract_text(result)

        assert "Server Procurement" in preview
        assert "$150,000" in preview

    def test_populate_handles_line_items(self, sample_xlsx):
        """Line items array should be inserted into spreadsheet."""
        data = {
            "line_items": [
                {"description": "Server", "quantity": 10, "unit_price": 5000, "total": 50000},
                {"description": "Storage", "quantity": 5, "unit_price": 10000, "total": 50000},
            ],
        }
        placeholder_map = {
            "line_items": "{{LINE_ITEMS}}",
        }

        result = XLSXPopulator.populate(sample_xlsx, data, placeholder_map)

        # Verify by loading result
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(result))
        ws = wb.active

        # Line items should be inserted starting at row 3
        assert ws["B3"].value == "Server"
        assert ws["B4"].value == "Storage"

    def test_extract_text_returns_markdown_table(self, sample_xlsx):
        """Text extraction should return markdown-formatted content."""
        preview = XLSXPopulator.extract_text(sample_xlsx)
        assert "IGCE" in preview  # Sheet title
        assert "|" in preview  # Table format


# ══════════════════════════════════════════════════════════════════════
#  Template Service Tests
# ══════════════════════════════════════════════════════════════════════


class TestTemplateService:
    """Tests for TemplateService."""

    @pytest.fixture
    def mock_generators(self):
        """Create mock markdown generators."""
        return {
            "sow": lambda title, data: f"# SOW: {title}\n\nGenerated content",
            "eval_criteria": lambda title, data: f"# Eval Criteria: {title}\n\nMarkdown only",
        }

    def test_markdown_only_uses_generator(self, mock_generators):
        """Markdown-only doc types should use generator directly."""
        service = TemplateService("tenant1", "user1", mock_generators)
        result = service.generate_document("eval_criteria", "Test Eval", {})

        assert result.success is True
        assert result.source == "markdown_fallback"
        assert result.file_type == "md"
        assert "Eval Criteria: Test Eval" in result.preview

    @patch("app.template_service._get_s3")
    def test_template_not_found_falls_back_to_markdown(self, mock_s3, mock_generators):
        """Missing S3 template should fall back to markdown."""
        from botocore.exceptions import ClientError

        mock_client = MagicMock()
        mock_client.get_object.side_effect = ClientError(
            {"Error": {"Code": "NoSuchKey"}}, "GetObject"
        )
        mock_s3.return_value = mock_client

        service = TemplateService("tenant1", "user1", mock_generators)
        result = service.generate_document("sow", "Test SOW", {})

        assert result.success is True
        assert result.source == "markdown_fallback"
        assert "SOW: Test SOW" in result.preview

    def test_no_generator_returns_error(self):
        """Missing generator should return error result."""
        service = TemplateService("tenant1", "user1", {})
        result = service._generate_markdown_fallback("unknown_type", "Test", {})

        assert result.success is False
        assert "No markdown generator" in result.error

    def test_template_result_dataclass(self):
        """TemplateResult should hold all expected fields."""
        result = TemplateResult(
            success=True,
            content=b"test content",
            preview="test preview",
            file_type="docx",
            source="s3_template",
            template_path="path/to/template.docx",
        )

        assert result.success is True
        assert result.content == b"test content"
        assert result.preview == "test preview"
        assert result.file_type == "docx"
        assert result.source == "s3_template"
        assert result.template_path == "path/to/template.docx"
        assert result.error is None


# ══════════════════════════════════════════════════════════════════════
#  Integration Tests
# ══════════════════════════════════════════════════════════════════════


class TestTemplateServiceIntegration:
    """Integration tests for template service with real DOCX/XLSX."""

    @pytest.fixture
    def real_generators(self):
        """Import actual generators from agentic_service."""
        try:
            from app.agentic_service import (
                _generate_acquisition_plan,
                _generate_cor_certification,
                _generate_contract_type_justification,
                _generate_eval_criteria,
                _generate_igce,
                _generate_justification,
                _generate_market_research,
                _generate_section_508,
                _generate_security_checklist,
                _generate_sow,
            )

            return {
                "sow": _generate_sow,
                "igce": _generate_igce,
                "market_research": _generate_market_research,
                "justification": _generate_justification,
                "acquisition_plan": _generate_acquisition_plan,
                "eval_criteria": _generate_eval_criteria,
                "security_checklist": _generate_security_checklist,
                "section_508": _generate_section_508,
                "cor_certification": _generate_cor_certification,
                "contract_type_justification": _generate_contract_type_justification,
            }
        except ImportError:
            pytest.skip("agentic_service generators not available")

    @patch("app.template_service._get_s3")
    def test_full_fallback_flow(self, mock_s3, real_generators):
        """Test complete fallback flow with real generators."""
        from botocore.exceptions import ClientError

        # Simulate S3 not having the template
        mock_client = MagicMock()
        mock_client.get_object.side_effect = ClientError(
            {"Error": {"Code": "NoSuchKey"}}, "GetObject"
        )
        mock_s3.return_value = mock_client

        service = TemplateService("test-tenant", "test-user", real_generators)

        # Test SOW generation with fallback
        result = service.generate_document(
            "sow",
            "Cloud Migration Services",
            {"description": "AWS migration", "period_of_performance": "12 months"},
        )

        assert result.success is True
        assert result.source == "markdown_fallback"
        assert "Cloud Migration Services" in result.preview
        assert "STATEMENT OF WORK" in result.preview

    @patch("app.template_service._get_s3")
    def test_igce_fallback_with_line_items(self, mock_s3, real_generators):
        """Test IGCE generation with line items data."""
        from botocore.exceptions import ClientError

        mock_client = MagicMock()
        mock_client.get_object.side_effect = ClientError(
            {"Error": {"Code": "NoSuchKey"}}, "GetObject"
        )
        mock_s3.return_value = mock_client

        service = TemplateService("test-tenant", "test-user", real_generators)

        result = service.generate_document(
            "igce",
            "Server Procurement IGCE",
            {
                "description": "Servers for data center",
                "total_estimate": "$500,000",
                "line_items": [
                    {"description": "Dell Servers", "quantity": 10, "unit_cost": "$25,000"},
                    {"description": "Storage Arrays", "quantity": 2, "unit_cost": "$100,000"},
                ],
            },
        )

        assert result.success is True
        assert "IGCE" in result.preview or "Cost Estimate" in result.preview
