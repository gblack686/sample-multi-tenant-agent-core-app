"""Template Service — Generates documents from S3 templates.

Fetches official NCI/HHS templates from S3, populates {{PLACEHOLDER}} tokens,
and returns the populated document. Falls back to markdown generation when
templates are unavailable.
"""
from __future__ import annotations

import io
import logging
import os
import re
import time
from dataclasses import dataclass
from typing import Any, Callable, Dict, List, Optional, Tuple

import boto3
from botocore.exceptions import BotoCoreError, ClientError

from app.template_registry import (
    TEMPLATE_BUCKET,
    get_alternate_s3_keys,
    get_placeholder_map,
    get_template_mapping,
    get_template_s3_key,
    has_template,
    is_markdown_only,
)

logger = logging.getLogger("eagle.template_service")

# ── Configuration ─────────────────────────────────────────────────────
CACHE_TTL_SECONDS = 60
S3_TIMEOUT_SECONDS = 10
DOCUMENTS_BUCKET = os.getenv("S3_BUCKET", "eagle-documents-695681773636-dev")

# ── S3 Singleton ──────────────────────────────────────────────────────
_s3_client = None


def _get_s3():
    global _s3_client
    if _s3_client is None:
        _s3_client = boto3.client("s3", region_name=os.getenv("AWS_REGION", "us-east-1"))
    return _s3_client


# ── Template Cache (60s TTL) ──────────────────────────────────────────
# Key: "{bucket}:{s3_key}"
# Value: {"ts": float, "data": bytes | None}
_template_cache: Dict[str, Dict] = {}
_MISS = object()


def _cache_key(bucket: str, s3_key: str) -> str:
    return f"{bucket}:{s3_key}"


def _cache_get(bucket: str, s3_key: str) -> object:
    """Return cached template bytes if fresh; _MISS sentinel if stale/absent."""
    entry = _template_cache.get(_cache_key(bucket, s3_key))
    if entry is not None and time.time() < entry["ts"] + CACHE_TTL_SECONDS:
        return entry["data"]
    return _MISS


def _cache_set(bucket: str, s3_key: str, data: Optional[bytes]) -> None:
    _template_cache[_cache_key(bucket, s3_key)] = {
        "ts": time.time(),
        "data": data,
    }


# ── Result Dataclass ──────────────────────────────────────────────────
@dataclass
class TemplateResult:
    """Result of template generation."""

    success: bool
    content: bytes  # Populated DOCX/XLSX bytes or markdown string encoded
    preview: str  # Markdown preview of content
    file_type: str  # "docx" | "xlsx" | "md"
    source: str  # "s3_template" | "markdown_fallback"
    template_path: Optional[str] = None  # S3 key of template used
    error: Optional[str] = None


# ══════════════════════════════════════════════════════════════════════
#  DOCX Populator
# ══════════════════════════════════════════════════════════════════════


class DOCXPopulator:
    """Populates DOCX templates by replacing {{PLACEHOLDER}} tokens."""

    @staticmethod
    def populate(
        template_bytes: bytes,
        data: Dict[str, Any],
        placeholder_map: Dict[str, str],
    ) -> bytes:
        """Replace placeholders in DOCX template with data values.

        Args:
            template_bytes: Raw DOCX file content
            data: Dict of field_name -> value
            placeholder_map: Dict of field_name -> {{PLACEHOLDER}}

        Returns:
            Populated DOCX bytes
        """
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed, cannot populate DOCX")
            raise ImportError("python-docx required for DOCX population")

        doc = Document(io.BytesIO(template_bytes))

        # Build replacement dict: {{PLACEHOLDER}} -> actual_value
        replacements = {}
        for field_name, placeholder in placeholder_map.items():
            if field_name in data:
                value = data[field_name]
                # Handle lists/arrays
                if isinstance(value, list):
                    if all(isinstance(v, str) for v in value):
                        value = "\n".join(f"- {v}" for v in value)
                    else:
                        # Assume list of dicts (line items)
                        lines = []
                        for item in value:
                            if isinstance(item, dict):
                                lines.append(", ".join(f"{k}: {v}" for k, v in item.items()))
                            else:
                                lines.append(str(item))
                        value = "\n".join(lines)
                replacements[placeholder] = str(value)

        # Replace in paragraphs
        for para in doc.paragraphs:
            DOCXPopulator._replace_in_paragraph(para, replacements)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        DOCXPopulator._replace_in_paragraph(para, replacements)

        # Replace in headers/footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        DOCXPopulator._replace_in_paragraph(para, replacements)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        DOCXPopulator._replace_in_paragraph(para, replacements)

        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    @staticmethod
    def _replace_in_paragraph(para, replacements: Dict[str, str]) -> None:
        """Replace placeholders in a paragraph while preserving formatting."""
        full_text = para.text
        for placeholder, value in replacements.items():
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, value)

        if full_text != para.text:
            # Clear and rebuild with new text (preserves first run's formatting)
            if para.runs:
                first_run = para.runs[0]
                for run in para.runs[1:]:
                    run.text = ""
                first_run.text = full_text
            else:
                para.text = full_text

    @staticmethod
    def extract_text(docx_bytes: bytes) -> str:
        """Extract text from DOCX as markdown preview."""
        try:
            from docx import Document
        except ImportError:
            return "[DOCX preview unavailable - python-docx not installed]"

        doc = Document(io.BytesIO(docx_bytes))
        lines = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            if "Heading 1" in style_name:
                lines.append(f"# {text}")
            elif "Heading 2" in style_name:
                lines.append(f"## {text}")
            elif "Heading 3" in style_name:
                lines.append(f"### {text}")
            else:
                lines.append(text)

        return "\n\n".join(lines)


# ══════════════════════════════════════════════════════════════════════
#  XLSX Populator
# ══════════════════════════════════════════════════════════════════════


class XLSXPopulator:
    """Populates XLSX templates by replacing {{PLACEHOLDER}} tokens."""

    @staticmethod
    def populate(
        template_bytes: bytes,
        data: Dict[str, Any],
        placeholder_map: Dict[str, str],
    ) -> bytes:
        """Replace placeholders in XLSX template with data values.

        Special handling for line_items: if {{LINE_ITEMS}} is found,
        expects data["line_items"] to be a list of dicts with keys like
        description, quantity, unit_price, total.

        Args:
            template_bytes: Raw XLSX file content
            data: Dict of field_name -> value
            placeholder_map: Dict of field_name -> {{PLACEHOLDER}}

        Returns:
            Populated XLSX bytes
        """
        try:
            from openpyxl import load_workbook
        except ImportError:
            logger.warning("openpyxl not installed, cannot populate XLSX")
            raise ImportError("openpyxl required for XLSX population")

        wb = load_workbook(io.BytesIO(template_bytes))

        # Build replacement dict
        replacements = {}
        line_items = None
        for field_name, placeholder in placeholder_map.items():
            if field_name in data:
                value = data[field_name]
                if field_name == "line_items" and isinstance(value, list):
                    line_items = value
                    replacements[placeholder] = "[LINE_ITEMS]"  # Marker
                else:
                    replacements[placeholder] = str(value)

        # Process all sheets
        for sheet in wb.worksheets:
            line_items_row = None

            for row_idx, row in enumerate(sheet.iter_rows(), start=1):
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        original = cell.value
                        for placeholder, value in replacements.items():
                            if placeholder in original:
                                if value == "[LINE_ITEMS]":
                                    line_items_row = row_idx
                                    cell.value = ""  # Clear placeholder
                                else:
                                    cell.value = original.replace(placeholder, value)
                                original = cell.value

            # Insert line items if found
            if line_items_row and line_items:
                XLSXPopulator._insert_line_items(sheet, line_items_row, line_items)

        # Save to bytes
        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()

    @staticmethod
    def _insert_line_items(sheet, start_row: int, items: List[Dict]) -> None:
        """Insert line items into spreadsheet starting at given row."""
        for i, item in enumerate(items):
            row = start_row + i
            # Standard IGCE columns: A=Item#, B=Description, C=Qty, D=Unit, E=UnitPrice, F=Total
            sheet.cell(row=row, column=1, value=i + 1)
            sheet.cell(row=row, column=2, value=item.get("description", ""))
            sheet.cell(row=row, column=3, value=item.get("quantity", 1))
            sheet.cell(row=row, column=4, value=item.get("unit", "EA"))
            sheet.cell(row=row, column=5, value=item.get("unit_price", 0))
            sheet.cell(row=row, column=6, value=item.get("total", item.get("unit_price", 0) * item.get("quantity", 1)))

    @staticmethod
    def extract_text(xlsx_bytes: bytes) -> str:
        """Extract text from XLSX as markdown preview."""
        try:
            from openpyxl import load_workbook
        except ImportError:
            return "[XLSX preview unavailable - openpyxl not installed]"

        wb = load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
        lines = []

        for sheet in wb.worksheets:
            lines.append(f"## {sheet.title}")
            lines.append("")

            for row in sheet.iter_rows(max_row=50):  # Limit preview
                cells = [str(cell.value) if cell.value is not None else "" for cell in row]
                if any(cells):
                    lines.append("| " + " | ".join(cells) + " |")

            lines.append("")

        return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════
#  Template Service
# ══════════════════════════════════════════════════════════════════════


class TemplateService:
    """Service for generating documents from S3 templates."""

    def __init__(
        self,
        tenant_id: str,
        user_id: str,
        markdown_generators: Optional[Dict[str, Callable]] = None,
    ):
        """Initialize template service.

        Args:
            tenant_id: Tenant identifier for S3 paths
            user_id: User identifier for S3 paths
            markdown_generators: Dict of doc_type -> generator function for fallback
        """
        self.tenant_id = tenant_id
        self.user_id = user_id
        self.markdown_generators = markdown_generators or {}

    def generate_document(
        self,
        doc_type: str,
        title: str,
        data: Dict[str, Any],
        output_format: str = "docx",
    ) -> TemplateResult:
        """Generate a document using S3 template or markdown fallback.

        Args:
            doc_type: Document type (sow, igce, etc.)
            title: Document title
            data: Data to populate template with
            output_format: Desired output format (docx, xlsx, md)

        Returns:
            TemplateResult with populated document
        """
        # Check if this doc_type should only use markdown
        if is_markdown_only(doc_type):
            return self._generate_markdown_fallback(doc_type, title, data)

        # Check if we have a template for this doc_type
        if not has_template(doc_type):
            logger.info("No template registered for %s, using markdown fallback", doc_type)
            return self._generate_markdown_fallback(doc_type, title, data)

        # Try to fetch and populate the template
        try:
            return self._generate_from_template(doc_type, title, data)
        except Exception as e:
            logger.warning(
                "Template generation failed for %s: %s, falling back to markdown",
                doc_type,
                str(e),
            )
            return self._generate_markdown_fallback(doc_type, title, data)

    def _generate_from_template(
        self,
        doc_type: str,
        title: str,
        data: Dict[str, Any],
    ) -> TemplateResult:
        """Generate document from S3 template."""
        mapping = get_template_mapping(doc_type)
        if not mapping:
            raise ValueError(f"No template mapping for {doc_type}")

        # Fetch template from S3 (with cache)
        template_bytes, s3_key = self._fetch_template(doc_type)
        if template_bytes is None:
            raise FileNotFoundError(f"Template not found for {doc_type}")

        # Add title to data
        data_with_title = {"title": title, **data}

        # Populate based on file type
        placeholder_map = get_placeholder_map(doc_type)
        if mapping.file_type == "docx":
            populated = DOCXPopulator.populate(template_bytes, data_with_title, placeholder_map)
            preview = DOCXPopulator.extract_text(populated)
            file_type = "docx"
        elif mapping.file_type == "xlsx":
            populated = XLSXPopulator.populate(template_bytes, data_with_title, placeholder_map)
            preview = XLSXPopulator.extract_text(populated)
            file_type = "xlsx"
        else:
            raise ValueError(f"Unknown file type: {mapping.file_type}")

        return TemplateResult(
            success=True,
            content=populated,
            preview=preview,
            file_type=file_type,
            source="s3_template",
            template_path=s3_key,
        )

    def _fetch_template(self, doc_type: str) -> Tuple[Optional[bytes], Optional[str]]:
        """Fetch template from S3, trying primary then alternates.

        Returns (template_bytes, s3_key) or (None, None) if not found.
        """
        s3 = _get_s3()
        bucket = TEMPLATE_BUCKET

        # Try primary template
        s3_key = get_template_s3_key(doc_type)
        if s3_key:
            template_bytes = self._fetch_s3_object(bucket, s3_key)
            if template_bytes:
                return template_bytes, s3_key

        # Try alternates
        for alt_key in get_alternate_s3_keys(doc_type):
            template_bytes = self._fetch_s3_object(bucket, alt_key)
            if template_bytes:
                logger.info("Using alternate template %s for %s", alt_key, doc_type)
                return template_bytes, alt_key

        return None, None

    def _fetch_s3_object(self, bucket: str, s3_key: str) -> Optional[bytes]:
        """Fetch object from S3 with caching."""
        # Check cache
        cached = _cache_get(bucket, s3_key)
        if cached is not _MISS:
            return cached

        # Fetch from S3
        try:
            s3 = _get_s3()
            response = s3.get_object(Bucket=bucket, Key=s3_key)
            data = response["Body"].read()
            _cache_set(bucket, s3_key, data)
            logger.debug("Fetched template from s3://%s/%s", bucket, s3_key)
            return data
        except ClientError as e:
            error_code = e.response.get("Error", {}).get("Code", "")
            if error_code in ("NoSuchKey", "404"):
                logger.debug("Template not found: s3://%s/%s", bucket, s3_key)
                _cache_set(bucket, s3_key, None)  # Cache the miss
            else:
                logger.warning("S3 error fetching %s: %s", s3_key, e)
            return None
        except BotoCoreError as e:
            logger.warning("S3 connection error fetching %s: %s", s3_key, e)
            return None
        except Exception as e:
            logger.error("Unexpected error fetching template %s: %s", s3_key, e)
            return None

    def _generate_markdown_fallback(
        self,
        doc_type: str,
        title: str,
        data: Dict[str, Any],
    ) -> TemplateResult:
        """Generate document using markdown generator."""
        generator = self.markdown_generators.get(doc_type)
        if not generator:
            return TemplateResult(
                success=False,
                content=b"",
                preview="",
                file_type="md",
                source="markdown_fallback",
                error=f"No markdown generator for {doc_type}",
            )

        try:
            markdown_content = generator(title, data)
            return TemplateResult(
                success=True,
                content=markdown_content.encode("utf-8"),
                preview=markdown_content,
                file_type="md",
                source="markdown_fallback",
            )
        except Exception as e:
            logger.error("Markdown generation failed for %s: %s", doc_type, e)
            return TemplateResult(
                success=False,
                content=b"",
                preview="",
                file_type="md",
                source="markdown_fallback",
                error=str(e),
            )


def create_template_service(
    tenant_id: str,
    user_id: str,
    markdown_generators: Optional[Dict[str, Callable]] = None,
) -> TemplateService:
    """Factory function to create a TemplateService instance."""
    return TemplateService(tenant_id, user_id, markdown_generators)
