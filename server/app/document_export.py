"""
Document Export Service
Generates DOCX and PDF files from markdown content.
"""
import io
import re
import logging
from datetime import datetime
from typing import Optional, Dict, Any

logger = logging.getLogger("eagle.export")

# ── DOCX Export ──────────────────────────────────────────────────────

def markdown_to_docx(content: str, title: str = "Document") -> bytes:
    """
    Convert markdown content to DOCX format.
    Returns the DOCX file as bytes.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        logger.warning("python-docx not installed, using fallback")
        return _docx_fallback(content, title)

    doc = Document()
    
    # Set document properties
    doc.core_properties.title = title
    doc.core_properties.author = "EAGLE - NCI Acquisition Assistant"
    doc.core_properties.created = datetime.utcnow()
    
    # Parse markdown and add to document
    lines = content.split('\n')
    in_table = False
    table_data = []
    in_code_block = False
    code_lines = []
    
    for line in lines:
        stripped = line.strip()
        
        # Handle code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                if code_lines:
                    p = doc.add_paragraph()
                    p.style = 'No Spacing'
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # Handle tables
        if stripped.startswith('|') and stripped.endswith('|'):
            if '---' in stripped:
                continue  # Skip separator row
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if not in_table:
                in_table = True
                table_data = [cells]
            else:
                table_data.append(cells)
            continue
        elif in_table:
            # End of table
            _add_table_to_doc(doc, table_data)
            table_data = []
            in_table = False
        
        # Handle headers
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
        # Handle horizontal rules
        elif stripped in ('---', '***', '___'):
            doc.add_paragraph('─' * 50)
        # Handle bullet points
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(stripped[2:], style='List Bullet')
        # Handle numbered lists
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(text, style='List Number')
        # Handle checkboxes
        elif stripped.startswith('- [ ] '):
            p = doc.add_paragraph('☐ ' + stripped[6:], style='List Bullet')
        elif stripped.startswith('- [x] ') or stripped.startswith('- [X] '):
            p = doc.add_paragraph('☑ ' + stripped[6:], style='List Bullet')
        # Handle blockquotes
        elif stripped.startswith('> '):
            p = doc.add_paragraph(stripped[2:])
            p.style = 'Quote'
        # Handle empty lines
        elif not stripped:
            doc.add_paragraph()
        # Handle regular paragraphs
        else:
            p = doc.add_paragraph()
            # Handle inline formatting
            _add_formatted_text(p, stripped)
    
    # Handle any remaining table
    if in_table and table_data:
        _add_table_to_doc(doc, table_data)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _add_table_to_doc(doc, table_data: list):
    """Add a table to the document."""
    if not table_data:
        return
    
    from docx.shared import Pt
    
    rows = len(table_data)
    cols = len(table_data[0]) if table_data else 0
    
    if rows == 0 or cols == 0:
        return
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                row.cells[j].text = cell_text
                # Bold header row
                if i == 0:
                    for paragraph in row.cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def _add_formatted_text(paragraph, text: str):
    """Add text with inline formatting (bold, italic, code) to a paragraph."""
    from docx.shared import Pt
    
    # Simple regex-based parsing for **bold**, *italic*, and `code`
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def _docx_fallback(content: str, title: str) -> bytes:
    """Fallback when python-docx is not available - return plain text."""
    header = f"EAGLE Document Export\nTitle: {title}\nGenerated: {datetime.utcnow().isoformat()}\n\n{'=' * 60}\n\n"
    return (header + content).encode('utf-8')


# ── PDF Export ───────────────────────────────────────────────────────

def markdown_to_pdf(content: str, title: str = "Document") -> bytes:
    """
    Convert markdown content to PDF format.
    Returns the PDF file as bytes.
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
    except ImportError:
        logger.warning("reportlab not installed, using fallback")
        return _pdf_fallback(content, title)
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    styles.add(ParagraphStyle(
        name='CustomTitle',
        parent=styles['Title'],
        fontSize=18,
        spaceAfter=20
    ))
    styles.add(ParagraphStyle(
        name='CustomHeading1',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=12,
        spaceBefore=20
    ))
    styles.add(ParagraphStyle(
        name='CustomHeading2',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=10,
        spaceBefore=16
    ))
    styles.add(ParagraphStyle(
        name='CustomHeading3',
        parent=styles['Heading3'],
        fontSize=12,
        spaceAfter=8,
        spaceBefore=12
    ))
    styles.add(ParagraphStyle(
        name='CustomCode',
        parent=styles['Code'],
        fontSize=8,
        fontName='Courier',
        backColor=colors.lightgrey
    ))
    styles.add(ParagraphStyle(
        name='CustomBullet',
        parent=styles['Normal'],
        leftIndent=20,
        bulletIndent=10
    ))
    
    story = []
    
    # Add title
    story.append(Paragraph(title, styles['CustomTitle']))
    story.append(Paragraph(f"Generated by EAGLE — {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Parse markdown
    lines = content.split('\n')
    in_code_block = False
    code_lines = []
    in_table = False
    table_data = []
    
    for line in lines:
        stripped = line.strip()
        
        # Handle code blocks
        if stripped.startswith('```'):
            if in_code_block:
                if code_lines:
                    code_text = '<br/>'.join(_escape_html(l) for l in code_lines)
                    story.append(Paragraph(code_text, styles['CustomCode']))
                    story.append(Spacer(1, 6))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # Handle tables
        if stripped.startswith('|') and stripped.endswith('|'):
            if '---' in stripped:
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if not in_table:
                in_table = True
                table_data = [cells]
            else:
                table_data.append(cells)
            continue
        elif in_table:
            _add_table_to_pdf(story, table_data, styles)
            table_data = []
            in_table = False
        
        # Handle headers
        if stripped.startswith('# '):
            story.append(Paragraph(_escape_html(stripped[2:]), styles['CustomHeading1']))
        elif stripped.startswith('## '):
            story.append(Paragraph(_escape_html(stripped[3:]), styles['CustomHeading2']))
        elif stripped.startswith('### '):
            story.append(Paragraph(_escape_html(stripped[4:]), styles['CustomHeading3']))
        elif stripped.startswith('#### '):
            story.append(Paragraph(_escape_html(stripped[5:]), styles['Heading4']))
        # Handle horizontal rules
        elif stripped in ('---', '***', '___'):
            story.append(Paragraph('─' * 60, styles['Normal']))
            story.append(Spacer(1, 6))
        # Handle bullet points
        elif stripped.startswith('- ') or stripped.startswith('* '):
            story.append(Paragraph(f"• {_escape_html(stripped[2:])}", styles['CustomBullet']))
        # Handle numbered lists
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            story.append(Paragraph(f"  {_escape_html(text)}", styles['CustomBullet']))
        # Handle empty lines
        elif not stripped:
            story.append(Spacer(1, 6))
        # Handle regular paragraphs
        else:
            # Convert markdown formatting to HTML-like tags for reportlab
            text = _markdown_to_reportlab(stripped)
            story.append(Paragraph(text, styles['Normal']))
    
    # Handle remaining table
    if in_table and table_data:
        _add_table_to_pdf(story, table_data, styles)
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def _add_table_to_pdf(story, table_data: list, styles):
    """Add a table to the PDF story."""
    from reportlab.platypus import Table, TableStyle, Paragraph
    from reportlab.lib import colors
    
    if not table_data:
        return
    
    # Wrap cell content in Paragraphs for word wrapping
    wrapped_data = []
    for row in table_data:
        wrapped_row = [Paragraph(_escape_html(cell), styles['Normal']) for cell in row]
        wrapped_data.append(wrapped_row)
    
    table = Table(wrapped_data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('TOPPADDING', (0, 1), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
    ]))
    
    story.append(table)
    from reportlab.platypus import Spacer
    story.append(Spacer(1, 12))


def _escape_html(text: str) -> str:
    """Escape HTML special characters."""
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;'))


def _markdown_to_reportlab(text: str) -> str:
    """Convert markdown inline formatting to reportlab tags."""
    # Escape HTML first
    text = _escape_html(text)
    # Bold
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    # Italic
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    # Code
    text = re.sub(r'`(.+?)`', r'<font name="Courier">\1</font>', text)
    return text


def _pdf_fallback(content: str, title: str) -> bytes:
    """Fallback when reportlab is not available - return plain text."""
    header = f"EAGLE Document Export (PDF fallback - text format)\nTitle: {title}\nGenerated: {datetime.utcnow().isoformat()}\n\n{'=' * 60}\n\n"
    return (header + content).encode('utf-8')


# ── Export Orchestrator ──────────────────────────────────────────────

def export_document(
    content: str,
    format: str,
    title: str = "Document",
    metadata: Optional[Dict[str, Any]] = None
) -> Dict[str, Any]:
    """
    Export document to specified format.
    
    Args:
        content: Markdown content
        format: "docx", "pdf", or "md"
        title: Document title
        metadata: Optional metadata to include
    
    Returns:
        Dict with: data (bytes), filename, content_type, size
    """
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    safe_title = re.sub(r'[^\w\-]', '_', title)[:50]
    
    if format.lower() == "docx":
        data = markdown_to_docx(content, title)
        filename = f"{safe_title}_{timestamp}.docx"
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif format.lower() == "pdf":
        data = markdown_to_pdf(content, title)
        filename = f"{safe_title}_{timestamp}.pdf"
        content_type = "application/pdf"
    elif format.lower() in ("md", "markdown"):
        data = content.encode('utf-8')
        filename = f"{safe_title}_{timestamp}.md"
        content_type = "text/markdown"
    else:
        raise ValueError(f"Unsupported format: {format}. Use docx, pdf, or md.")
    
    return {
        "data": data,
        "filename": filename,
        "content_type": content_type,
        "size_bytes": len(data),
        "format": format.lower(),
        "generated_at": datetime.utcnow().isoformat(),
    }
